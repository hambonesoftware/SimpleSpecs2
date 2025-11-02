"""Services for persisting specification approvals and exports."""

from __future__ import annotations

import csv
import hashlib
import io
import json
import zipfile
from datetime import UTC, datetime, timedelta
from pathlib import Path
from typing import Any, Iterable, Tuple

from docx import Document as DocxDocument
from sqlalchemy import asc
from sqlmodel import Session, select

from ..config import Settings
from ..middleware import get_request_id
from ..models import Document, SpecAuditEntry, SpecRecord


class SpecRecordError(RuntimeError):
    """Base error for spec record operations."""


def ensure_document(session: Session, document_id: int) -> Document:
    """Return the document for the given id or raise an error."""

    document = session.get(Document, document_id)
    if document is None:
        raise SpecRecordError(f"Document {document_id} not found")
    return document


def _serialise_payload(payload: Any) -> dict[str, Any]:
    """Return a JSON-serialisable payload copy."""

    try:
        serialised = json.loads(json.dumps(payload, ensure_ascii=False))
    except (TypeError, ValueError) as exc:  # pragma: no cover - defensive branch
        raise SpecRecordError("Payload is not JSON serialisable") from exc
    if not isinstance(serialised, dict):
        raise SpecRecordError("Payload must be a JSON object")
    return serialised


def _payload_hash(payload: dict[str, Any]) -> str:
    """Return a stable hash for the payload."""

    canonical = json.dumps(
        payload, sort_keys=True, ensure_ascii=False, separators=(",", ":")
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def approve_specifications(
    session: Session,
    *,
    document: Document,
    payload: Any,
    reviewer: str,
    notes: str | None = None,
) -> SpecRecord:
    """Create or update the frozen specification payload for a document."""

    if document.id is None:
        raise SpecRecordError("Document must be persisted before approval")

    data = _serialise_payload(payload)
    payload_hash = _payload_hash(data)
    now = datetime.now(UTC)

    record_stmt = select(SpecRecord).where(SpecRecord.document_id == document.id)
    record = session.exec(record_stmt).first()

    detail: dict[str, Any] = {"content_hash": payload_hash}
    request_id = get_request_id()
    if request_id:
        detail["request_id"] = request_id
    if notes:
        detail["notes"] = notes

    if record is None:
        record = SpecRecord(
            document_id=document.id,
            state="approved",
            reviewer=reviewer,
            created_at=now,
            updated_at=now,
            approved_at=now,
            frozen_at=now,
            content_hash=payload_hash,
            payload=data,
        )
        session.add(record)
        session.flush()
        summary = "Specifications approved and frozen"
        detail["changed"] = True
    else:
        changed = record.content_hash != payload_hash or record.state != "approved"
        record.updated_at = now
        if changed:
            detail["previous_hash"] = record.content_hash
            detail["changed"] = True
            record.payload = data
            record.content_hash = payload_hash
            record.state = "approved"
            record.approved_at = now
            record.frozen_at = now
            record.reviewer = reviewer
            summary = "Specifications re-approved with updates"
        else:
            detail["changed"] = False
            if not record.reviewer:
                record.reviewer = reviewer
            summary = "Approval replayed without changes"
        session.add(record)
        session.flush()

    audit_detail = detail.copy()
    if record.id is None:
        session.refresh(record)
    if record.id is None:
        raise SpecRecordError("Specification record failed to persist")

    audit = SpecAuditEntry(
        document_id=document.id,
        record_id=record.id,
        action="approve",
        actor=reviewer,
        summary=summary,
        detail=audit_detail,
    )
    session.add(audit)
    session.commit()
    session.refresh(record)
    return record


def fetch_spec_record(
    session: Session, *, document_id: int
) -> Tuple[SpecRecord | None, list[SpecAuditEntry]]:
    """Return the spec record and audit entries for a document."""

    record_stmt = select(SpecRecord).where(SpecRecord.document_id == document_id)
    record = session.exec(record_stmt).first()

    audit_stmt = (
        select(SpecAuditEntry)
        .where(SpecAuditEntry.document_id == document_id)
        .order_by(
            asc(SpecAuditEntry.__table__.c.created_at)  # type: ignore[attr-defined]
        )
    )
    audit_entries = list(session.exec(audit_stmt))
    return record, audit_entries


def export_spec_record(
    session: Session,
    *,
    record: SpecRecord,
    settings: Settings,
    fmt: str,
    actor: str | None = None,
) -> tuple[Path, str]:
    """Generate an export artifact for a frozen specification record."""

    export_dir = settings.export_dir
    export_dir.mkdir(parents=True, exist_ok=True)
    _cleanup_exports(export_dir, settings.export_retention_days)

    suffix = (record.content_hash or "unhashed")[:12]
    request_id = get_request_id()
    if fmt == "csv":
        path = export_dir / f"spec-{record.document_id}-{suffix}.zip"
        media_type = "application/zip"
        _write_csv_bundle(record, path)
    elif fmt == "docx":
        path = export_dir / f"spec-{record.document_id}-{suffix}.docx"
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        _write_docx_report(record, path)
    else:
        raise SpecRecordError(f"Unsupported export format: {fmt}")

    detail = {"format": fmt, "filename": path.name}
    if request_id:
        detail["request_id"] = request_id

    audit = SpecAuditEntry(
        document_id=record.document_id,
        record_id=record.id,
        action="export",
        actor=actor,
        summary=f"Specification export generated ({fmt})",
        detail=detail,
    )
    session.add(audit)
    session.commit()
    return path, media_type


def _cleanup_exports(export_dir: Path, retention_days: int) -> None:
    """Remove export artifacts older than the retention window."""

    if retention_days <= 0:
        return

    cutoff = datetime.now(UTC) - timedelta(days=retention_days)
    for item in export_dir.glob("*"):
        if not item.is_file():
            continue
        try:
            modified = datetime.fromtimestamp(item.stat().st_mtime, UTC)
        except OSError:  # pragma: no cover - filesystem race
            continue
        if modified < cutoff:
            try:
                item.unlink()
            except OSError:  # pragma: no cover - filesystem race
                continue


def _iter_bucket_items(
    record: SpecRecord,
) -> Iterable[tuple[str, list[dict[str, Any]]]]:
    payload = record.payload or {}
    buckets = payload.get("buckets", {})
    if not isinstance(buckets, dict):
        return []
    return sorted(
        (
            (discipline, list(lines) if isinstance(lines, list) else [])
            for discipline, lines in buckets.items()
        ),
        key=lambda item: item[0],
    )


def _write_csv_bundle(record: SpecRecord, target: Path) -> None:
    """Create a ZIP archive containing per-discipline CSV files."""

    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as bundle:
        for discipline, lines in _iter_bucket_items(record):
            buffer = io.StringIO()
            writer = csv.DictWriter(
                buffer,
                fieldnames=["discipline", "text", "page", "header", "source"],
            )
            writer.writeheader()
            for line in lines:
                page = line.get("page")
                page_value = page + 1 if isinstance(page, int) else ""
                header_path = line.get("header_path")
                if isinstance(header_path, list):
                    header_value = " > ".join(str(segment) for segment in header_path)
                else:
                    header_value = ""
                writer.writerow(
                    {
                        "discipline": discipline,
                        "text": str(line.get("text", "")),
                        "page": page_value,
                        "header": header_value,
                        "source": str(line.get("source", "")),
                    }
                )
            bundle.writestr(f"{discipline}.csv", buffer.getvalue())


def _write_docx_report(record: SpecRecord, target: Path) -> None:
    """Create a DOCX report with header hierarchy for each discipline."""

    document = DocxDocument()
    document.add_heading(f"Specifications for Document {record.document_id}", level=1)

    meta_lines = [
        f"State: {record.state}",
        f"Reviewer: {record.reviewer or '—'}",
    ]
    if record.approved_at:
        meta_lines.append(f"Approved: {record.approved_at.isoformat()}")
    document.add_paragraph(" | ".join(meta_lines))

    for discipline, lines in _iter_bucket_items(record):
        document.add_heading(_humanise_discipline(discipline), level=2)
        if not lines:
            document.add_paragraph("No specification lines in this bucket.")
            continue
        for line in lines:
            paragraph = document.add_paragraph(str(line.get("text", "")))
            page = line.get("page")
            page_value = page + 1 if isinstance(page, int) else "—"
            header_path = line.get("header_path")
            if isinstance(header_path, list) and header_path:
                header_value = " > ".join(str(segment) for segment in header_path)
            else:
                header_value = "No header context"
            meta_run = paragraph.add_run(
                f"\nPage {page_value} • Header: {header_value}"
            )
            meta_run.italic = True

    document.save(target)


def _humanise_discipline(value: str) -> str:
    return value.replace("_", " ").title()


__all__ = [
    "SpecRecordError",
    "approve_specifications",
    "ensure_document",
    "export_spec_record",
    "fetch_spec_record",
]
